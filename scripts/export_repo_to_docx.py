from __future__ import annotations

from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "Chipset-Defect-Vision_full_workspace_export.docx"
SKIP_DIRS = {".git"}
CODE_DIR_HINTS = {
    "app",
    "frontend",
    "scripts",
    "training",
}
CODE_FILE_NAMES = {
    "dockerfile",
    "makefile",
    "readme.md",
    "requirements.txt",
    ".gitignore",
    ".dockerignore",
    "data.yaml",
    "classes.txt",
}
CODE_EXTENSIONS = {
    ".bat",
    ".c",
    ".cc",
    ".cfg",
    ".conf",
    ".css",
    ".csv",
    ".env",
    ".go",
    ".graphql",
    ".h",
    ".hpp",
    ".html",
    ".ini",
    ".java",
    ".js",
    ".json",
    ".jsx",
    ".md",
    ".mjs",
    ".properties",
    ".ps1",
    ".py",
    ".rb",
    ".scss",
    ".sh",
    ".sql",
    ".svg",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}
NON_CODE_EXTENSIONS = {
    ".cache",
    ".doc",
    ".docx",
    ".gif",
    ".jpeg",
    ".jpg",
    ".pdf",
    ".pickle",
    ".png",
    ".pth",
    ".pt",
    ".pyc",
    ".webp",
    ".zip",
}
NON_CODE_PATH_PARTS = {
    "__pycache__",
    "data/labels",
    "data/images",
    "incoming_data",
    "raw_data",
    "runs",
    "training/data",
    "weights",
}
MAX_TEXT_FILE_SIZE = 512 * 1024


def relative_posix(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def iter_dirs_and_files() -> tuple[list[Path], list[Path]]:
    dirs: list[Path] = []
    files: list[Path] = []
    for current_root, dir_names, file_names in os_walk_sorted(ROOT):
        current_path = Path(current_root)
        dirs.append(current_path)
        for file_name in file_names:
            files.append(current_path / file_name)
    return dirs, files


def os_walk_sorted(root: Path) -> Iterable[tuple[str, list[str], list[str]]]:
    import os

    for current_root, dir_names, file_names in os.walk(root):
        dir_names[:] = sorted(name for name in dir_names if name not in SKIP_DIRS)
        file_names.sort()
        yield current_root, dir_names, file_names


def looks_like_code_file(path: Path) -> bool:
    rel = relative_posix(path)
    lower_rel = rel.lower()
    if path.name.lower() == OUTPUT_PATH.name.lower():
        return False
    if path.suffix.lower() in NON_CODE_EXTENSIONS:
        return False
    if any(part in lower_rel for part in NON_CODE_PATH_PARTS):
        return False
    if path.name.lower() in CODE_FILE_NAMES:
        return is_text_file(path)
    if path.suffix.lower() in CODE_EXTENSIONS:
        return is_text_file(path)
    if any(part in CODE_DIR_HINTS for part in path.parts):
        return is_text_file(path)
    return False


def is_text_file(path: Path) -> bool:
    try:
        if path.stat().st_size > MAX_TEXT_FILE_SIZE:
            return False
        raw = path.read_bytes()
    except OSError:
        return False
    if b"\x00" in raw:
        return False
    try:
        raw.decode("utf-8")
        return True
    except UnicodeDecodeError:
        try:
            raw.decode("latin-1")
            return True
        except UnicodeDecodeError:
            return False


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return "[Unable to decode file content]"


def set_code_style(document: Document) -> None:
    styles = document.styles
    if "CodeBlock" in styles:
        return
    style = styles.add_style("CodeBlock", 1)
    style.base_style = styles["Normal"]
    style.font.name = "Consolas"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    style.font.size = Pt(8)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CodeBlock")
    run = paragraph.add_run(text if text else "[Empty file]")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)


def build_document() -> tuple[Path, int, int, int]:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    set_code_style(document)

    dirs, files = iter_dirs_and_files()
    code_files = [path for path in files if looks_like_code_file(path)]
    non_code_files = [path for path in files if path not in code_files]
    extension_counts = Counter(path.suffix.lower() or "[no extension]" for path in files)

    document.add_heading("Chipset Defect Vision Workspace Export", level=0)
    document.add_paragraph(f"Root folder: {ROOT}")
    document.add_paragraph(f"Total folders (excluding .git): {len(dirs)}")
    document.add_paragraph(f"Total files (excluding .git): {len(files)}")
    document.add_paragraph(f"Code/text files embedded: {len(code_files)}")
    document.add_paragraph(f"Non-code files listed by name: {len(non_code_files)}")

    document.add_heading("Extension Summary", level=1)
    for extension, count in sorted(extension_counts.items()):
        document.add_paragraph(f"{extension}: {count}", style="List Bullet")

    document.add_heading("Folder Inventory", level=1)
    for directory in dirs:
        rel = "." if directory == ROOT else relative_posix(directory)
        document.add_paragraph(rel, style="List Bullet")

    document.add_heading("All File Names", level=1)
    for file_path in files:
        rel = relative_posix(file_path)
        label = "[code]" if file_path in code_files else "[name only]"
        document.add_paragraph(f"{label} {rel}", style="List Bullet")

    document.add_page_break()
    document.add_heading("Embedded Code And Text Files", level=1)
    for index, file_path in enumerate(code_files, start=1):
        rel = relative_posix(file_path)
        document.add_heading(f"{index}. {rel}", level=2)
        document.add_paragraph(f"Size: {file_path.stat().st_size} bytes")
        add_code_block(document, read_text_file(file_path))
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH, len(dirs), len(files), len(code_files)


if __name__ == "__main__":
    output_path, folder_count, file_count, code_count = build_document()
    print(f"Created: {output_path}")
    print(f"Folders: {folder_count}")
    print(f"Files: {file_count}")
    print(f"Embedded code/text files: {code_count}")
